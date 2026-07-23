# -*- coding: utf-8 -*-
"""Revise the SRS docx to match the implemented CoffeeShop Lite system.

Produces Group2_1.SRS_v2.docx next to the original. All edits are documented
in SRS_CHANGELOG.md. The original file is never modified.
"""
import copy
import shutil
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph

DOCS = Path(r"d:\STUDY\Tai_lieu\FPT_University\Semester\Summer_2026\SWP391\Documents")
SRC = DOCS / "Group2_1.SRS Document.docx"
DST = DOCS / "Group2_1.SRS_v2.docx"
SHOTS = DOCS / "srs_screenshots"

MOBILE_W = Inches(2.9)
DESKTOP_W = Inches(5.9)

log_lines = []


def log(msg):
    log_lines.append(msg)
    print(msg.encode("ascii", "replace").decode("ascii"))


# ---------------------------------------------------------------- helpers

def all_paragraphs(doc):
    for p in doc.element.body.iter(qn("w:p")):
        yield Paragraph(p, doc)


def all_tables(doc):
    return [Table(t, doc) for t in doc.element.body.iter(qn("w:tbl"))]


def find_para(doc, text, exact=False):
    for p in all_paragraphs(doc):
        t = p.text.strip()
        if (exact and t == text) or (not exact and text in t):
            return p
    return None


def clear_runs(para):
    for r in list(para._p.findall(qn("w:r"))):
        para._p.remove(r)


def set_para_lines(para, lines):
    """Replace paragraph content with lines separated by soft breaks."""
    clear_runs(para)
    for i, line in enumerate(lines):
        run = para.add_run(line)
        if i < len(lines) - 1:
            run._r.append(run._r.makeelement(qn("w:br"), {}))


def replace_para(doc, old, new, must=True):
    p = find_para(doc, old)
    if p is None:
        if must:
            log(f"!! NOT FOUND: {old[:80]}")
        return None
    set_para_lines(p, [new] if isinstance(new, str) else new)
    log(f"edited: {old[:60]} -> {str(new)[:60]}")
    return p


def delete_para(doc, text):
    p = find_para(doc, text)
    if p is None:
        log(f"!! NOT FOUND (delete): {text[:80]}")
        return
    p._p.getparent().remove(p._p)
    log(f"deleted: {text[:70]}")


def set_cell(cell, value):
    """value: str or list of lines."""
    ps = cell.paragraphs
    for extra in ps[1:]:
        extra._p.getparent().remove(extra._p)
    set_para_lines(ps[0], [value] if isinstance(value, str) else value)


def clone_row(table, template_idx, texts, insert_after_idx=None):
    tr = copy.deepcopy(table.rows[template_idx]._tr)
    if insert_after_idx is None:
        table._tbl.append(tr)
    else:
        table.rows[insert_after_idx]._tr.addnext(tr)
    from docx.table import _Row
    row = _Row(tr, table)
    for cell, text in zip(row.cells, texts):
        set_cell(cell, text)
    return row


def insert_paragraph_after(doc, anchor_p, text="", style=None, italic=False, bold=False):
    new_p = copy.deepcopy(anchor_p._p) if False else None
    para = Paragraph(anchor_p._p.makeelement(qn("w:p"), {}), doc)
    anchor_p._p.addnext(para._p)
    if style:
        try:
            para.style = style
        except Exception:
            pass
    if text:
        run = para.add_run(text)
        run.italic = italic
        run.bold = bold
    return para


def insert_image_after(doc, anchor_p, img_path, width, caption):
    """Insert image paragraph + caption after anchor. Returns caption para."""
    img_para = insert_paragraph_after(doc, anchor_p)
    img_para.alignment = 1
    run = img_para.add_run()
    run.add_picture(str(img_path), width=width)
    cap = insert_paragraph_after(doc, img_para, caption, italic=True)
    cap.alignment = 1
    for r in cap.runs:
        r.font.size = Pt(10)
    return cap


def section_image_anchor(doc, heading_text, max_scan=16):
    """Find the drawing paragraph in a section, or the heading itself."""
    h = find_para(doc, heading_text)
    if h is None:
        log(f"!! heading not found: {heading_text}")
        return None
    el = h._p
    for _ in range(max_scan):
        el = el.getnext()
        if el is None:
            break
        if el.tag == qn("w:p"):
            p = Paragraph(el, doc)
            style = p.style.name if p.style else ""
            if style.startswith("Heading") and p.text.strip():
                break
            if el.findall(".//" + qn("w:drawing")):
                return p
    return h


# ---------------------------------------------------------------- steps

def step_catalog(doc, tables):
    """Table 1.3.1: remove barista split rows, fix UC-52, add missing UCs, renumber."""
    cat = tables[2]

    def row_uc(r):
        return r.cells[1].text.strip().split("\n")[0]

    # remove barista split rows (26, 27)
    removed = 0
    for r in list(cat.rows):
        uc = row_uc(r)
        if r.cells[0].text.strip().isdigit() and uc.startswith(("Split Bill by Item", "Create New Bill")) \
                and "Barista" in r.cells[3].text or (uc.startswith("Create New Bill") and "separated out" in r.cells[3].text):
            if uc.startswith("Split Bill by Item") or "separated out" in r.cells[3].text:
                r._tr.getparent().remove(r._tr)
                removed += 1
    log(f"catalog: removed {removed} barista split rows")

    # update cashier Split Bill description
    for r in cat.rows:
        if row_uc(r) == "Split Bill" and "Cashier" in r.cells[3].text:
            set_cell(r.cells[3], "Cashier splits a single order into multiple bills so guests can pay "
                                 "separately. At least one item must remain on the original bill; the new "
                                 "bill is created in Served status and is locked from further splitting.")
            log("catalog: updated cashier Split Bill description")

    # UC Monitor Staff Screen -> Access Staff Screens
    for r in cat.rows:
        if row_uc(r) == "Monitor Staff Screen":
            set_cell(r.cells[1], "Access Staff Screens")
            set_cell(r.cells[3], "Manager can open any staff role screen (Barista, Cashier, Waiter) from the "
                                 "admin session to operate or verify staff workflows. No separate live "
                                 "monitoring feature is provided.")
            log("catalog: renamed Monitor Staff Screen -> Access Staff Screens")

    # insert barista item-level preparation row after "Mark Order Ready"
    for i, r in enumerate(cat.rows):
        if row_uc(r) == "Mark Order Ready":
            clone_row(cat, i, ["0", "Prepare Item Line", "Order Preparation",
                               "Barista marks each item line of an order as prepared (per-item preparation). "
                               "When every line is fully prepared the order becomes Ready automatically."],
                      insert_after_idx=i)
            log("catalog: added Prepare Item Line (barista)")
            break

    # insert manager rows after "Manage Staff & Shifts"
    mgr_new = [
        ["0", "View Monthly Payroll", "Staff Management",
         "Manager views total worked shifts and hours per staff for a selected month. "
         "Evening shift counts 5 hours, other shifts 6 hours; only completed shifts are counted."],
        ["0", "Manage Ingredient Inventory", "Inventory Management",
         "Manager manages ingredient stock, unit, minimum stock, and import cost. The system warns on "
         "low stock and automatically disables menu items whose ingredients run out."],
        ["0", "Manage Recipes", "Inventory Management",
         "Manager defines the ingredient recipe of each menu item; recipes drive automatic inventory "
         "deduction when the Barista finishes preparing an order."],
        ["0", "Import Menu from Excel", "Configuration",
         "Manager downloads a generated Excel template and imports menu items in bulk from a filled "
         "Excel file; invalid rows are skipped and reported."],
    ]
    for i, r in enumerate(cat.rows):
        if row_uc(r).startswith("Manage Staff & Shifts"):
            after = i
            for texts in reversed(mgr_new):
                clone_row(cat, i, texts, insert_after_idx=after)
            log("catalog: added 4 manager rows (payroll, inventory, recipes, excel import)")
            break

    # renumber
    counter = 0
    for r in cat.rows:
        v = r.cells[0].text.strip()
        if v.isdigit() or v == "0":
            counter += 1
            set_cell(r.cells[0], str(counter))
    log(f"catalog: renumbered 1..{counter}")


def step_business_rules(doc, tables):
    br = tables[38]
    # delete BR-01 (operating hours)
    for r in list(br.rows):
        if r.cells[0].text.strip() == "BR-01":
            r._tr.getparent().remove(r._tr)
            log("BR: deleted BR-01 (operating hours - not implemented)")
            break
    # renumber BR-02..BR-12 -> BR-01..BR-11
    n = 0
    for r in br.rows:
        if r.cells[0].text.strip().startswith("BR-"):
            n += 1
            set_cell(r.cells[0], f"BR-{n:02d}")
    # rewrite payment rule (old BR-03, now BR-02)
    for r in br.rows:
        txt = r.cells[1].text
        if "confirm payment" in txt:
            set_cell(r.cells[1], "Cashier can confirm payment only for an order in Served status. Payment "
                                 "is confirmed per order; other active orders of the same table are paid "
                                 "separately.")
            log("BR: rewrote payment rule to per-order (matches code)")
        if "Cup stock decreases" in txt:
            set_cell(r.cells[1], "Cup stock decreases by drink quantity when Barista marks an order Ready. "
                                 "If the remaining cup stock is insufficient, the system blocks the Ready "
                                 "transition until Admin restocks. Food/cake items do not reduce cup stock.")
            log("BR: cup rule updated to hard block")
    # append new rules
    last = len(br.rows) - 1
    clone_row(br, last, [f"BR-{n + 1:02d}",
                         "Only Cashier or Admin can split a bill. At least one item must remain on the "
                         "original order; the split-off order is created in Served status and both orders "
                         "are locked from repeated splitting."])
    clone_row(br, last, [f"BR-{n + 2:02d}",
                         "When an order finishes preparation, linked recipe ingredients are deducted from "
                         "inventory. Menu items with insufficient ingredients are disabled automatically "
                         "and new orders exceeding the remaining servings are rejected."])
    log(f"BR: appended BR-{n+1:02d} (split bill) and BR-{n+2:02d} (inventory)")


def step_traceability(doc, tables):
    tr = tables[15]
    remap = {f"BR-{i:02d}": f"BR-{i-1:02d}" for i in range(2, 13)}
    for r in tr.rows[1:]:
        cell = r.cells[1]
        txt = cell.text
        refs = [x.strip() for x in txt.replace("\n", " ").split(",") if x.strip()]
        new_refs = []
        for ref in refs:
            if ref == "BR-01":
                continue  # operating hours removed
            new_refs.append(remap.get(ref, ref))
        uc = r.cells[0].text.strip()
        if uc == "UC-06" and "BR-12" not in new_refs:
            new_refs.append("BR-12")  # split rule
        if uc == "UC-04" and "BR-13" not in new_refs:
            new_refs.append("BR-13")  # ingredient deduction
        set_cell(cell, ", ".join(new_refs))
    last = len(tr.rows) - 1
    clone_row(tr, last, ["UC-09", "BR-07", "MSG02, MSG16"])
    clone_row(tr, last + 1, ["UC-10", "BR-07, BR-13", "MSG01, MSG02, MSG16"])
    clone_row(tr, last + 2, ["UC-11", "BR-07", "MSG02, MSG16"])
    log("traceability: BR references renumbered, UC-09..11 rows added")


def step_coverage(doc, tables):
    cov = tables[6]
    last = len(cov.rows) - 1
    clone_row(cov, last, ["UC-09", "Manage Staff, Shifts & Payroll", "Staff Management", "Manager"])
    clone_row(cov, last + 1, ["UC-10", "Manage Inventory & Recipes", "Inventory Management", "Manager"])
    clone_row(cov, last + 2, ["UC-11", "Import Menu from Excel", "Administration", "Manager"])
    log("coverage summary: UC-09..UC-11 added")


def step_entities(doc, tables):
    ent = tables[5]
    for r in ent.rows:
        name = r.cells[1].text.strip()
        if name == "Orders":
            cur = r.cells[2].text.strip()
            set_cell(r.cells[2], cur + " Also stores split-lock and invoice-printed flags (splitLocked, "
                                       "invoicePrinted) and a reserved customerPhone field that is currently "
                                       "always empty.")
        if name == "OrderItems":
            cur = r.cells[2].text.strip()
            set_cell(r.cells[2], cur + " Includes preparedQty for item-level preparation tracking.")
    last = len(ent.rows) - 1
    base = int(ent.rows[last].cells[0].text.strip() or 10)
    rows = [
        ["Staff", "Stores individual staff records: name, role, personal PIN, username/password, and "
                  "active status. Used by staff management and payroll. In the current version personal "
                  "PINs are not used for staff login (login uses shared role accounts in Users)."],
        ["Shifts", "Stores shift assignments per staff member and date (Morning, Afternoon, Evening), "
                   "assigned role, working hours, status, and notes. A unique constraint plus overlap "
                   "checking prevents double assignment of the same staff/shift/date."],
        ["Inventory", "Stores ingredient stock with unit, minimum stock level, and import cost. Drives "
                      "low-stock dashboard warnings and automatic disabling of out-of-stock menu items."],
        ["RecipeItems", "Stores the ingredient quantities required by each menu item. Used to deduct "
                        "inventory automatically when the Barista finishes preparing an order and to "
                        "compute the remaining sellable servings."],
    ]
    for i, (name, desc) in enumerate(rows):
        clone_row(ent, last, [str(base + 1 + i), name, desc])
    log("entities: added Staff, Shifts, Inventory, RecipeItems; Orders/OrderItems columns noted")

    # note under the ERD image
    anchor = section_image_anchor(doc, "1.5 Entity Relationship Diagram")
    if anchor is not None:
        insert_paragraph_after(
            doc, anchor,
            "Note: the diagram above shows the core entities. The implementation additionally uses the "
            "Staff, Shifts, Inventory, and RecipeItems tables described below (13 tables in total).",
            italic=True)
        log("entities: ERD note inserted")


def step_spec_edits(doc):
    # --- UC-02 A2 (order status without table context)
    replace_para(
        doc,
        "A2 - Customer opens status page without table context",
        ["A1 - Table has no active orders: System displays an empty state.",
         "A2 - Customer opens status page without table context: System suggests scanning the table QR; "
         "a manual order-number lookup is available but limited to orders created in the current "
         "browsing session.",
         "A3 - Order has moved to a final cleaned state: System no longer shows it as active customer order."])

    # --- UC-06 preconditions + alternative flows (per-order payment, split)
    replace_para(
        doc,
        "For table orders, all related items have been served by Waiter",
        ["- Cashier has logged in with a valid PIN.",
         "- The selected order has been served by Waiter (payment is checked per order).",
         "- The bill is still unpaid.",
         "- The order total is calculated from item price, size extra price, and quantity."])
    replace_para(
        doc,
        "A1 - Table still has unserved orders: System prevents payment.",
        ["A1 - Selected order is not yet served: System prevents payment. Payment is validated per "
         "order; other orders of the same table are paid separately.",
         "A2 - Customer orders at the counter: Cashier opens Counter Order, selects table/items/notes, "
         "and confirms with hold confirmation.",
         "A3 - Database error occurs: System keeps the bill unpaid and shows an error message.",
         "A4 - Guests want to pay separately: Cashier opens Split Bill and moves selected quantities to "
         "a new bill. At least one item must remain on the original bill; the new bill is created in "
         "Served status and both bills are locked from repeated splitting."])

    # --- UC-04 flow + A3 cup block + item prepare
    replace_para(
        doc,
        "3. Barista holds a Pending order for 0.5 second to move it to Preparing.",
        ["1. Barista opens the preparation board.",
         "2. System shows Pending, Preparing, and Ready filters/columns.",
         "3. Barista holds a Pending order for 0.5 second to move it to Preparing.",
         "4. Barista prepares the items and reads notes; each finished item line can be marked prepared "
         "individually (item-level preparation).",
         "5. Barista holds a Preparing order for 0.5 second to move it to Ready, or the order becomes "
         "Ready automatically when every item line is prepared.",
         "6. System deducts cups for drink quantities, deducts recipe ingredients from inventory, and "
         "notifies Waiter."])
    replace_para(
        doc,
        "A3 - Cup stock is low: System keeps showing cup count",
        ["A1 - Barista tries to move Ready back to Preparing/Pending: System blocks the backward transition.",
         "A2 - Order has only food/cake items: Cup stock is not deducted.",
         "A3 - Cup stock is insufficient for the drinks in the order: System blocks the Preparing-to-Ready "
         "transition and shows the required versus remaining cup count; Admin must restock before the "
         "order can be marked Ready."])

    # --- 3.2.1 remove non-existent search
    delete_para(doc, "Search orders by order number or table name.")

    # --- 3.2.2 per-order payment guard
    replace_para(
        doc,
        "Prevents payment if the table still has unserved orders.",
        "Prevents payment unless the selected order is in Served status (payment is validated per order).")
    replace_para(
        doc,
        "Confirm that all ordered items have already been served by Waiter.",
        "Confirm that the selected order has already been served by Waiter.")

    # --- 3.3.2 item-level prepare + inventory bullets
    p = find_para(doc, "Deduct cup stock by drink item quantity after marking drink orders Ready.")
    if p is not None:
        p2 = insert_paragraph_after(
            doc, p, "Mark each item line as prepared; the order becomes Ready automatically when every "
                    "line is fully prepared (item-level preparation).")
        p2._p.set(qn("w:dummy"), "") if False else None
        insert_paragraph_after(
            doc, p2, "Trigger automatic deduction of linked recipe ingredients from inventory when "
                     "preparation completes; menu items with insufficient ingredients are disabled "
                     "automatically.")
        try:
            p2.style = p.style
        except Exception:
            pass
        log("3.3.2: added item-prepare and inventory bullets")
    p = find_para(doc, "Prevents invalid backward status transitions.")
    if p is not None:
        insert_paragraph_after(
            doc, p, "Blocks the Ready transition when cup stock or ingredient stock is insufficient.")
        log("3.3.2: added cup/ingredient block bullet")

    # --- 3.1.3 session-limited manual lookup note
    p = find_para(doc, "Avoid entering an order number manually when the table context already exists.")
    if p is not None:
        insert_paragraph_after(
            doc, p, "Use the manual order-number lookup only for orders created in the current browsing "
                    "session (guest lookup is session-limited).")
        log("3.1.3: session-limited lookup bullet added")

    # --- 3.5.5 PIN reality
    replace_para(
        doc,
        "Each role has a PIN/password stored in the Users table.",
        "Barista, Cashier, and Waiter role PINs are stored in the Users table "
        "(demo values: Barista 1111, Cashier 2222, Waiter/Runner 3333; Admin account 8888).")
    replace_para(
        doc,
        "Change demo PIN values in the database when required by the project owner.",
        "Change staff demo PIN values in the Users table when required. Note: the admin dashboard PIN "
        "(8888) is currently fixed in the application code and cannot be changed from the database.")
    replace_para(
        doc,
        "The system has predefined roles: Admin, Barista, Cashier, and Waiter.",
        "The system has predefined roles: Admin, Barista, Cashier, and Waiter (the Waiter role is "
        "implemented under the internal role key 'runner').")

    # --- 4.1.3 database contents
    replace_para(
        doc,
        "The database shall store users, tables, menu items, item sizes, orders, order items, cash events",
        "The database shall store users, tables, menu items, item sizes, orders, order items, cash "
        "events, store state, system logs, staff records, shift assignments, ingredient inventory, and "
        "menu item recipes.")

    # --- actors table naming
    tables = all_tables(doc)
    for r in tables[1].rows:
        if r.cells[1].text.strip() == "Waiter":
            set_cell(r.cells[1], "Waiter (Runner)")
            log("actors: Waiter -> Waiter (Runner)")

    # --- note under barista / cashier use case diagrams
    for head, note in [
        ("1.3.2.3 UCs for Barista",
         "Note: bill splitting is performed by the Cashier (or Admin) only; the Barista has no split-bill "
         "function in the implemented system."),
        ("1.3.2.4 UCs for Cashier",
         "Note: Split Bill is restricted to Cashier/Admin. At least one item must remain on the original "
         "bill and the split-off bill is created in Served status."),
    ]:
        anchor = section_image_anchor(doc, head)
        if anchor is not None:
            insert_paragraph_after(doc, anchor, note, italic=True)
            log(f"diagram note added: {head}")


def step_new_sections_2x(doc, tables):
    """Insert 2.7 Staff/Shifts/Payroll and 2.8 Inventory specs, rename 2.7 -> 2.9."""
    trace_head = find_para(doc, "2.7 Traceability Matrix")
    if trace_head is None:
        log("!! 2.7 Traceability Matrix heading not found")
        return
    set_para_lines(trace_head, ["2.9 Traceability Matrix"])

    spec_template = tables[7]._tbl  # UC-01 spec table

    def make_spec(uc_id, actors, secondary, desc, pre, post, flow, alt):
        tbl_el = copy.deepcopy(spec_template)
        t = Table(tbl_el, doc)
        vals = [uc_id, actors, secondary, desc, pre, post, flow, alt]
        labels = ["Use Case ID", "Primary Actors", "Secondary Actors", "Description",
                  "Preconditions", "Postconditions", "Normal Sequence/Flow",
                  "Alternative Sequences/Flows"]
        for row, label, val in zip(t.rows, labels, vals):
            set_cell(row.cells[0], label)
            set_cell(row.cells[1], val)
        return tbl_el

    # build blocks in reverse order and addprevious before the traceability heading
    anchor_el = trace_head._p

    def add_heading_before(text, style):
        p = Paragraph(anchor_el.makeelement(qn("w:p"), {}), doc)
        anchor_el.addprevious(p._p)
        try:
            p.style = style
        except Exception:
            pass
        p.add_run(text)
        return p

    def add_table_before(tbl_el):
        anchor_el.addprevious(tbl_el)
        spacer = Paragraph(anchor_el.makeelement(qn("w:p"), {}), doc)
        anchor_el.addprevious(spacer._p)

    add_heading_before("2.7 Staff, Shifts and Payroll", "Heading 2")
    add_heading_before("2.7.1 Manage Staff, Shifts and Payroll", "Heading 3")
    add_table_before(make_spec(
        "UC-09", "Manager", "Barista, Cashier, Waiter",
        "As a manager, I want to manage staff records, assign work shifts without overlaps, and view "
        "monthly worked hours so that staffing and payroll are transparent.",
        ["- Manager has unlocked the admin dashboard with the admin PIN.",
         "- Staff data exists or can be created."],
        ["- Staff records are created/updated/deactivated with per-person PIN and role.",
         "- Shift assignments are stored per staff, date, and shift without overlaps.",
         "- Monthly payroll shows total shifts and hours per staff member."],
        ["1. Manager opens the Staff Management screen.",
         "2. Manager creates or edits a staff member with name, role, PIN, and status.",
         "3. Manager assigns the staff member to a shift (Morning 06:00-12:00, Afternoon 12:00-18:00, "
         "Evening 18:00-23:00) on a selected date with an assigned role.",
         "4. System validates that the staff member has no overlapping assignment for the same "
         "date and shift.",
         "5. Manager marks completed shifts (worked/finished status).",
         "6. Manager opens the payroll section, selects a month, and reviews total shifts and hours "
         "per staff member."],
        ["A1 - Overlapping shift assignment: System rejects the assignment (SHIFT_OVERLAP).",
         "A2 - Shift not completed: Shifts without a completed status are excluded from payroll totals.",
         "A3 - Payroll hours: Evening shift counts 5 hours, other shifts count 6 hours."]))

    add_heading_before("2.8 Inventory and Menu Data", "Heading 2")
    add_heading_before("2.8.1 Manage Inventory and Recipes", "Heading 3")
    add_table_before(make_spec(
        "UC-10", "Manager", "Barista, Customer / Guest",
        "As a manager, I want to manage ingredient stock and menu item recipes so that ingredient usage "
        "is tracked automatically and out-of-stock items stop being sold.",
        ["- Manager has unlocked the admin dashboard.",
         "- Ingredients and recipes exist or can be created."],
        ["- Ingredient stock, unit, minimum level, and import cost are stored.",
         "- Each menu item may define a recipe of ingredient quantities.",
         "- Inventory is deducted automatically when the Barista finishes preparing an order.",
         "- Menu items with insufficient ingredients are disabled automatically.",
         "- Low-stock warnings are shown on the admin dashboard."],
        ["1. Manager opens the Inventory screen.",
         "2. Manager creates or edits ingredients (name, unit, stock, minimum stock, import cost).",
         "3. Manager assigns recipe quantities to menu items in Menu Management.",
         "4. When Barista completes preparation, the system deducts the linked ingredient quantities.",
         "5. The system disables menu items whose ingredients are exhausted and warns when stock "
         "falls below the minimum level."],
        ["A1 - Order exceeds remaining servings: System rejects order creation.",
         "A2 - Ingredient deleted or stock adjusted: System refreshes menu availability immediately.",
         "A3 - Stock below minimum: Dashboard shows a low-stock warning listing the ingredients."]))

    add_heading_before("2.8.2 Import Menu from Excel", "Heading 3")
    add_table_before(make_spec(
        "UC-11", "Manager", "-",
        "As a manager, I want to import menu items in bulk from an Excel file so that a large menu can "
        "be set up quickly.",
        ["- Manager has unlocked the admin dashboard.",
         "- Manager has a filled Excel file following the provided template."],
        ["- Valid rows are imported as menu items; invalid rows are skipped.",
         "- The import result (imported/skipped counts) is reported and logged."],
        ["1. Manager opens Menu Management and downloads the generated Excel template.",
         "2. Manager fills product rows (bilingual names, category, price, sizes, image path).",
         "3. Manager uploads the Excel file through the Import function.",
         "4. System validates each row, imports valid products, and skips invalid rows.",
         "5. System shows the number of imported and skipped rows and records a system log."],
        ["A1 - File format invalid: System rejects the file with an error message.",
         "A2 - Some rows invalid: System imports the valid rows and reports the skipped count."]))

    log("section 2: added 2.7 (UC-09), 2.8 (UC-10, UC-11); traceability renamed to 2.9")


def step_new_sections_3x(doc):
    """Add 3.8 Staff/Shift/Payroll and 3.9 Inventory screens before section 4."""
    sec4 = find_para(doc, "4. Non-Functional Requirements", exact=True)
    if sec4 is None:
        log("!! section 4 heading not found")
        return
    anchor_el = sec4._p

    def add_before(text, style=None, italic=False):
        p = Paragraph(anchor_el.makeelement(qn("w:p"), {}), doc)
        anchor_el.addprevious(p._p)
        if style:
            try:
                p.style = style
            except Exception:
                pass
        if text:
            run = p.add_run(text)
            run.italic = italic
        return p

    def add_image_before(img, width, caption):
        p = add_before("")
        p.alignment = 1
        p.add_run().add_picture(str(SHOTS / img), width=width)
        cap = add_before(caption, italic=True)
        cap.alignment = 1
        for r in cap.runs:
            r.font.size = Pt(10)

    add_before("3.8 Staff, Shift and Payroll Management", "Heading 3")
    add_before("3.8.1 Staff Management Screen", "Heading 4")
    add_before("Platform: Desktop Web / Tablet Web")
    add_before("Primary Actor: Manager")
    add_before("Related Use Case: UC-09 - Manage Staff, Shifts and Payroll")
    add_image_before("17-admin-staff.png", DESKTOP_W,
                     "Implemented screen: staff list, weekly shift board, and shift assignment form")
    add_before("This screen allows the Manager to:")
    for b in [
        "View, add, edit, and deactivate staff members with name, role, and personal PIN.",
        "Assign staff to Morning (06:00-12:00), Afternoon (12:00-18:00), or Evening (18:00-23:00) "
        "shifts per day and per role column.",
        "See missing-staff warnings for unassigned shift slots on the weekly board.",
        "Update shift status (scheduled, worked, finished, absent).",
        "Be prevented from assigning the same staff member to overlapping shifts (SHIFT_OVERLAP).",
    ]:
        add_before(b)
    add_before("3.8.2 Monthly Payroll Section", "Heading 4")
    add_image_before("18-admin-payroll.png", DESKTOP_W,
                     "Implemented screen: monthly worked-hours summary (payroll) per staff member")
    add_before("This section allows the Manager to:")
    for b in [
        "Select a month and view total completed shifts and total hours per staff member.",
        "Filter the payroll summary by role.",
        "Rely on the counting rule: only shifts with worked/finished status are counted; the Evening "
        "shift counts 5 hours and other shifts count 6 hours.",
    ]:
        add_before(b)

    add_before("3.9 Inventory and Recipe Management", "Heading 3")
    add_before("3.9.1 Inventory Screen", "Heading 4")
    add_before("Platform: Desktop Web / Tablet Web")
    add_before("Primary Actor: Manager")
    add_before("Related Use Case: UC-10 - Manage Inventory and Recipes")
    add_image_before("19-admin-inventory.png", DESKTOP_W,
                     "Implemented screen: ingredient inventory with stock levels and low-stock warnings")
    add_before("This screen allows the Manager to:")
    for b in [
        "View, add, edit, and delete ingredients with unit, stock, minimum stock, and import cost.",
        "See which menu items are automatically disabled when their ingredients run out.",
        "Get low-stock warnings on the admin dashboard when stock falls below the minimum level.",
        "Maintain menu item recipes (ingredient quantities) through Menu Management so that "
        "preparation deducts inventory automatically.",
    ]:
        add_before(b)
    log("section 3: added 3.8 Staff/Shift/Payroll and 3.9 Inventory screens")


def step_screenshots(doc):
    mapping = [
        ("3.1.1 Customer Web Menu Screen", "02-menu-customer.png", MOBILE_W,
         "Implemented screen: customer mobile menu after QR table detection"),
        ("3.1.2 Order Summary / Checkout Screen", "03-menu-cart.png", MOBILE_W,
         "Implemented screen: cart summary and hold-to-confirm checkout"),
        ("3.1.3 Table Order Tracking Function", "04-order-status.png", MOBILE_W,
         "Implemented screen: customer order status for the locked table"),
        ("3.2.1 POS Order List Screen", "10-cashier-unpaid.png", DESKTOP_W,
         "Implemented screen: cashier unpaid/paid order tabs with cash strip"),
        ("3.2.2 Payment Processing Function", "11-cashier-split.png", DESKTOP_W,
         "Implemented screen: split-bill dialog (Cashier moves selected quantities to a new bill)"),
        ("3.2.3 Counter Order Function", "12-counter-order.png", DESKTOP_W,
         "Implemented screen: counter order for walk-in customers"),
        ("3.3.1 KDS Order Board Screen", "06-barista-board.png", DESKTOP_W,
         "Implemented screen: barista preparation board with cup stock and per-order/per-item modes"),
        ("3.4.1 Wait Station / Table Layout Screen", "08-runner-station.png", DESKTOP_W,
         "Implemented screen: waiter service station with table layout"),
        ("3.4.2 Move / Merge Table Function", "09-table-transfer.png", DESKTOP_W,
         "Implemented screen: table transfer"),
        ("3.5.1 Admin Dashboard Screen", "14-admin-dashboard.png", DESKTOP_W,
         "Implemented screen: admin dashboard with revenue, low-stock warning, and table map"),
        ("3.5.2 Product and Menu Management Screen", "15-admin-menu.png", DESKTOP_W,
         "Implemented screen: menu management including Excel import/template buttons"),
        ("3.5.3 Table and QR Management Screen", "16-admin-tables.png", DESKTOP_W,
         "Implemented screen: tables and QR management"),
        ("3.5.4 Reports Dashboard Screen", "20-system-logs.png", DESKTOP_W,
         "Implemented screen: bilingual system logs used for reporting and audit"),
        ("3.5.5 Role Access and Demo Staff Accounts", "13-admin-pin-gate.png", DESKTOP_W,
         "Implemented screen: admin PIN overlay protecting the dashboard"),
        ("3.6.1 Staff PIN Login Screen", "05-staff-login.png", DESKTOP_W,
         "Implemented screen: staff role selection and PIN login"),
    ]
    for heading, img, width, caption in mapping:
        anchor = section_image_anchor(doc, heading)
        if anchor is None:
            continue
        path = SHOTS / img
        if not path.exists():
            log(f"!! screenshot missing: {img}")
            continue
        insert_image_after(doc, anchor, path, width, caption)
        log(f"screenshot embedded: {heading} <- {img}")


def step_nfr(doc):
    # missing 4.2 heading
    p421 = find_para(doc, "4.2.1 Usability", exact=True)
    if p421 is not None:
        h = Paragraph(p421._p.makeelement(qn("w:p"), {}), doc)
        p421._p.addprevious(h._p)
        try:
            h.style = "Heading 3"
        except Exception:
            pass
        h.add_run("4.2 Quality Attributes")
        log("NFR: inserted missing '4.2 Quality Attributes' heading")

    # stray Heading-4 bullet
    stray = find_para(doc, "The table QR mechanism should support more tables")
    if stray is not None:
        try:
            stray.style = "normal"
        except Exception:
            try:
                stray.style = "Normal"
            except Exception:
                pass
        log("NFR: fixed stray Heading style on QR scalability bullet")

    # known limitations under 4.2.3 Security
    p = find_para(doc, "System logs shall be accessible only to Manager/Admin for auditing.")
    if p is not None:
        intro = insert_paragraph_after(
            doc, p, "Known limitations (accepted in the current student demo, to be addressed in future "
                    "versions):", bold=True)
        bullets = [
            "GET /api/tables is public and returns each table's access code, allowing the customer menu "
            "to be opened without physically scanning the table QR.",
            "Staff passwords and PINs are stored and transferred in plaintext; the admin staff API "
            "returns them in its JSON response for admin sessions.",
            "The API allows cross-origin requests from any origin (Access-Control-Allow-Origin: *).",
            "Revenue reports aggregate by order creation time (createdAt); no separate payment "
            "timestamp is stored, so an order paid on a later day is counted on its creation day.",
        ]
        anchor = intro
        for b in bullets:
            anchor = insert_paragraph_after(doc, anchor, "- " + b)
        log("NFR: known-limitations block added to 4.2.3 Security")


def step_record_of_changes(doc, tables):
    roc = tables[0]
    clone_row(roc, len(roc.rows) - 1, [
        "21/7", "A, M, D", "Group 2",
        "Synced SRS with the implemented system: split bill assigned to Cashier only; payment rule "
        "changed to per-order; operating-hours rule removed and business rules renumbered; added "
        "staff/shift/payroll (UC-09), inventory/recipes (UC-10), and Excel menu import (UC-11); ERD "
        "extended to 13 entities; item-level preparation and cup-stock hard block documented; "
        "implemented screen captures embedded; security limitations documented."])
    log("record of changes: new row added")


def step_toc_refresh(doc):
    settings = doc.settings.element
    if settings.find(qn("w:updateFields")) is None:
        el = settings.makeelement(qn("w:updateFields"), {qn("w:val"): "true"})
        settings.append(el)
    log("TOC: updateFields flag set (Word will refresh page numbers on open)")


def main():
    shutil.copyfile(SRC, DST)
    doc = Document(str(DST))
    tables = all_tables(doc)

    step_catalog(doc, tables)
    step_business_rules(doc, tables)
    step_traceability(doc, tables)
    step_coverage(doc, tables)
    step_entities(doc, tables)
    step_spec_edits(doc)
    step_new_sections_2x(doc, tables)
    step_new_sections_3x(doc)
    step_screenshots(doc)
    step_nfr(doc)
    step_record_of_changes(doc, tables)
    step_toc_refresh(doc)

    doc.save(str(DST))
    (DOCS / "SRS_CHANGELOG_worklog.txt").write_text("\n".join(log_lines), encoding="utf-8")
    print(f"\nSaved {DST}")


if __name__ == "__main__":
    main()
