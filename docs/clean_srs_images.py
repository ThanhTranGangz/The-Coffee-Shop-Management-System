# -*- coding: utf-8 -*-
"""Replace FR section screenshots in SRS_v2: remove old mockups + wrong embeds,
insert exactly one correct implemented screenshot per section.
"""
import shutil
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph

DOCS = Path(r"d:\STUDY\Tai_lieu\FPT_University\Semester\Summer_2026\SWP391\Documents")
SRC = DOCS / "Group2_1.SRS_v2.docx"
DST = DOCS / "Group2_1.SRS_v2.docx"
SHOTS = DOCS / "srs_screenshots"
BACKUP = DOCS / "Group2_1.SRS_v2.before_image_clean.docx"

MOBILE_W = Inches(2.9)
DESKTOP_W = Inches(5.8)

# One correct screenshot per FR screen section
SECTION_IMAGES = [
    ("3.1.1 Customer Web Menu Screen", "02-menu-customer.png", MOBILE_W,
     "Màn hình thực tế: menu khách sau khi quét QR / mở bằng tableCode"),
    ("3.1.2 Order Summary / Checkout Screen", "03-menu-cart.png", MOBILE_W,
     "Màn hình thực tế: giỏ hàng (sửa số lượng, bỏ món, tổng tiền, ghi chú)"),
    ("3.1.3 Table Order Tracking Function", "04-order-status.png", MOBILE_W,
     "Màn hình thực tế: theo dõi trạng thái đơn của bàn / phiên hiện tại"),
    ("3.2.1 POS Order List Screen", "10-cashier-unpaid.png", DESKTOP_W,
     "Màn hình thực tế: thu ngân — tab chưa thanh toán / đã thanh toán (không có ô tìm kiếm)"),
    ("3.2.2 Payment Processing Function", "11-cashier-split.png", DESKTOP_W,
     "Màn hình thực tế: dialog tách hóa đơn trên màn thu ngân (thanh toán từng đơn Served; tách bill khi khách trả riêng)"),
    ("3.2.3 Counter Order Function", "12-counter-order.png", DESKTOP_W,
     "Màn hình thực tế: gọi món tại quầy (Cashier/Admin chọn bàn và món)"),
    ("3.3.1 KDS Order Board Screen", "06-barista-board.png", DESKTOP_W,
     "Màn hình thực tế: bảng pha chế theo đơn (Pending / Preparing / Ready) và số cốc còn"),
    ("3.3.2 Update Item Status Function", "07-barista-item-prepare.png", DESKTOP_W,
     "Màn hình thực tế: chế độ Nấu theo món — đánh dấu preparedQty từng dòng món"),
    ("3.4.1 Wait Station / Table Layout Screen", "08-runner-station.png", DESKTOP_W,
     "Màn hình thực tế: bồi bàn (Waiter/Runner) — phục vụ, dọn bàn, sơ đồ bàn"),
    ("3.4.2 Move / Merge Table Function", "09-table-transfer.png", DESKTOP_W,
     "Màn hình thực tế: đổi bàn (chuyển toàn bộ đơn đang mở sang bàn đích)"),
    ("3.5.1 Admin Dashboard Screen", "14-admin-dashboard.png", DESKTOP_W,
     "Màn hình thực tế: dashboard admin — doanh thu, bán chạy, cảnh báo low-stock, sơ đồ bàn"),
    ("3.5.2 Product and Menu Management Screen", "15-admin-menu.png", DESKTOP_W,
     "Màn hình thực tế: quản lý thực đơn (CRUD, công thức món, Tải file mẫu / Import Excel)"),
    ("3.5.3 Table and QR Management Screen", "16-admin-tables.png", DESKTOP_W,
     "Màn hình thực tế: quản lý bàn và mã QR"),
    ("3.5.4 Reports Dashboard Screen", "14b-admin-reports.png", DESKTOP_W,
     "Màn hình thực tế: báo cáo doanh thu trên dashboard (lọc ngày/tuần/tháng/năm) — không phải trang log"),
    ("3.5.5 Role Access and Demo Staff Accounts", "13-admin-pin-gate.png", DESKTOP_W,
     "Màn hình thực tế: cổng PIN admin (8888) trước khi vào dashboard"),
    ("3.6.1 Staff PIN Login Screen", "05-staff-login.png", DESKTOP_W,
     "Màn hình thực tế: đăng nhập PIN theo role (Barista 1111 / Cashier 2222 / Waiter-Runner 3333)"),
    ("3.8.1 Staff Management Screen", "17-admin-staff.png", DESKTOP_W,
     "Màn hình thực tế: quản lý nhân viên và phân ca (chống trùng ca)"),
    ("3.8.2 Monthly Payroll Section", "18-admin-payroll.png", DESKTOP_W,
     "Màn hình thực tế: bảng chấm công tháng (Ca Tối 5h, ca khác 6h; chỉ ca Đã làm/Hoàn thành)"),
    ("3.9.1 Inventory Screen", "19-admin-inventory.png", DESKTOP_W,
     "Màn hình thực tế: kho nguyên liệu (tồn kho, min stock, cảnh báo)"),
]

SYSTEM_LOGS_AFTER = "Review system logs for customer, barista, cashier, waiter, and admin actions."
SYSTEM_LOGS_IMG = ("20-system-logs.png", DESKTOP_W,
                   "Màn hình thực tế: Log hệ thống (lọc theo actor/role) — trang riêng system-logs.jsp")


def log(msg):
    print(msg.encode("ascii", "replace").decode("ascii"))


def find_para(doc, text, exact=False):
    for p in doc.element.body.iter(qn("w:p")):
        para = Paragraph(p, doc)
        t = para.text.strip()
        if (exact and t == text) or (not exact and text in t):
            return para
    return None


def next_siblings_until_heading(start_p):
    """Yield following paragraph elements until next Heading* with text."""
    el = start_p._p.getnext()
    while el is not None:
        if el.tag == qn("w:p"):
            p = Paragraph(el, start_p._parent)
            style = (p.style.name if p.style else "") or ""
            text = p.text.strip()
            if style.startswith("Heading") and text:
                break
            yield el, p
        elif el.tag == qn("w:tbl"):
            yield el, None
        elif el.tag == qn("w:sdt"):
            yield el, None
        el = el.getnext()


def is_caption(text):
    t = (text or "").strip()
    return (
        t.startswith("Implemented screen")
        or t.startswith("Màn hình thực tế")
        or t.startswith("Màn hình:")
    )


def remove_images_under_heading(doc, heading_text):
    h = find_para(doc, heading_text, exact=True)
    if h is None:
        h = find_para(doc, heading_text, exact=False)
    if h is None:
        log(f"!! heading missing: {heading_text}")
        return None
    removed = 0
    for el, p in list(next_siblings_until_heading(h)):
        if p is None:
            continue
        has_drawing = bool(el.findall(".//" + qn("w:drawing")))
        if has_drawing or is_caption(p.text):
            el.getparent().remove(el)
            removed += 1
    log(f"cleared {removed} image/caption para(s) under [{heading_text}]")
    return h


def insert_image_after(doc, anchor_p, img_path, width, caption):
    # image paragraph
    img_p = Paragraph(anchor_p._p.makeelement(qn("w:p"), {}), doc)
    anchor_p._p.addnext(img_p._p)
    img_p.alignment = 1
    run = img_p.add_run()
    run.add_picture(str(img_path), width=width)
    # caption
    cap_p = Paragraph(img_p._p.makeelement(qn("w:p"), {}), doc)
    img_p._p.addnext(cap_p._p)
    cap_p.alignment = 1
    r = cap_p.add_run(caption)
    r.italic = True
    r.font.size = Pt(10)
    return cap_p


def find_insert_anchor(heading_p):
    """Prefer after 'Related Use Case' / first IMG-slot; else after heading."""
    anchor = heading_p
    for el, p in next_siblings_until_heading(heading_p):
        if p is None:
            continue
        t = p.text.strip()
        if t.startswith("Related Use Case") or t.startswith("Related Use Cases"):
            return p
        if t.startswith("Platform:") or t.startswith("Primary Actor"):
            anchor = p
        # stop before long body text starting with "This screen"
        if t.startswith("This screen") or t.startswith("This function") or t.startswith("This section"):
            return anchor
    return anchor


def ensure_system_logs_section(doc):
    """Add 3.5.6 System Logs after 3.5.5 block (before 3.6), with correct screenshot."""
    if find_para(doc, "3.5.6 System Logs Screen", exact=True):
        # refresh image only
        h = remove_images_under_heading(doc, "3.5.6 System Logs Screen")
        if h is not None:
            img, w, cap = SYSTEM_LOGS_IMG
            insert_image_after(doc, find_insert_anchor(h) or h, SHOTS / img, w, cap)
        return

    h55 = find_para(doc, "3.5.5 Role Access and Demo Staff Accounts", exact=True)
    h36 = find_para(doc, "3.6 Security and Staff Identification", exact=True)
    if h36 is None:
        log("!! cannot insert 3.5.6 — 3.6 missing")
        return
    # Insert headings/content before 3.6
    def add_before(text, style=None):
        p = Paragraph(h36._p.makeelement(qn("w:p"), {}), doc)
        h36._p.addprevious(p._p)
        if style:
            try:
                p.style = style
            except Exception:
                pass
        if text:
            p.add_run(text)
        return p

    # build in reverse order so final order is correct
    img, w, cap = SYSTEM_LOGS_IMG
    # we'll insert heading then content then image after heading
    # reverse: last inserted is closest to 3.6
    bullets = [
        "Filter logs by actor / role (Customer, Admin, Barista, Cashier, Waiter).",
        "Refresh the bilingual audit trail of login, order, payment, inventory, and admin actions.",
        "Use logs for presentation and troubleshooting (read-only for Admin).",
    ]
    for b in reversed(bullets):
        add_before(b)
    add_before("This screen allows the Manager to:")
    # placeholder for image — insert after Related Use Case line
    add_before("Related Use Case: UC-08 - Dashboard, Cash, Cups & Logs")
    add_before("Primary Actor: Manager")
    add_before("Platform: Desktop Web / Tablet Web")
    h = add_before("3.5.6 System Logs Screen", "Heading 4")
    # now find Related Use Case we just added and insert image after it
    # walk forward from h
    for el, p in next_siblings_until_heading(h):
        if p and p.text.strip().startswith("Related Use Case"):
            insert_image_after(doc, p, SHOTS / img, w, cap)
            break
    log("added section 3.5.6 System Logs Screen")


def fix_reports_section_text(doc):
    """Clarify 3.5.4 is the dashboard revenue/report view, not system logs."""
    p = find_para(doc, "View sales reports by selected date range.")
    # nothing mandatory; caption already clarifies
    # Fix any leftover wrong caption text if present as plain para
    for el in list(doc.element.body.iter(qn("w:p"))):
        para = Paragraph(el, doc)
        t = para.text.strip()
        if "system-logs" in t.lower() and "3.5.4" in t:
            pass
        if t.startswith("Implemented screen: bilingual system logs"):
            el.getparent().remove(el)
            log("removed stale system-logs caption under wrong section")


def add_split_bullet_if_missing(doc):
    p = find_para(doc, "View unpaid served bills that are ready for payment.")
    if p is None:
        return
    # check if split bullet already exists nearby
    found = False
    for el, q in next_siblings_until_heading(find_para(doc, "3.2.1 POS Order List Screen", exact=True) or p):
        if q and "split" in q.text.lower() or (q and "tách" in q.text.lower()):
            found = True
            break
    if not found:
        new_p = Paragraph(p._p.makeelement(qn("w:p"), {}), doc)
        p._p.addnext(new_p._p)
        new_p.add_run(
            "Open Split Bill on a served unpaid order to move selected item quantities to a new bill "
            "(must leave at least one item on the original bill)."
        )
        try:
            new_p.style = p.style
        except Exception:
            pass
        log("added split-bill bullet under 3.2.1")


def add_item_prepare_note_under_332(doc):
    h = find_para(doc, "3.3.2 Update Item Status Function", exact=True)
    if h is None:
        return
    # ensure image was placed; text already updated in prior revision


def main():
    shutil.copyfile(SRC, BACKUP)
    doc = Document(str(SRC))

    for heading, img, width, caption in SECTION_IMAGES:
        h = remove_images_under_heading(doc, heading)
        if h is None:
            continue
        path = SHOTS / img
        if not path.exists():
            log(f"!! missing file {img}")
            continue
        anchor = find_insert_anchor(h)
        insert_image_after(doc, anchor, path, width, caption)
        log(f"set [{heading}] <- {img}")

    # 3.3.2 previously had no image in original; ensure it has item-prepare
    # (already in SECTION_IMAGES)

    fix_reports_section_text(doc)
    add_split_bullet_if_missing(doc)
    ensure_system_logs_section(doc)

    # Also strip any orphaned "Implemented screen" captions still floating
    orphans = 0
    for el in list(doc.element.body.iter(qn("w:p"))):
        para = Paragraph(el, doc)
        if is_caption(para.text) and not el.findall(".//" + qn("w:drawing")):
            # keep captions that immediately follow a drawing — check previous sibling
            prev = el.getprevious()
            if prev is not None and prev.tag == qn("w:p") and prev.findall(".//" + qn("w:drawing")):
                # rewrite English "Implemented" captions already set by us — keep
                continue
            # orphan caption without preceding image
            if para.text.strip().startswith("Implemented screen"):
                el.getparent().remove(el)
                orphans += 1
    log(f"removed {orphans} orphan Implemented captions")

    doc.save(str(DST))
    log(f"saved {DST}")
    log(f"backup {BACKUP}")


if __name__ == "__main__":
    main()
