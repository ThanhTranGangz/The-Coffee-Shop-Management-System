# -*- coding: utf-8 -*-
"""Fix SRS issues: cashier image, no-UI sentences, 3.3.2/3.3.3 split, UC diagram styles."""
import shutil
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph

DOCS = Path(r"d:\STUDY\Tai_lieu\FPT_University\Semester\Summer_2026\SWP391\Documents")
SRC = DOCS / "Group2_1.SRS_v2.docx"
BACKUP = DOCS / "Group2_1.SRS_v2.before_layout_fix.docx"
SHOTS = DOCS / "srs_screenshots"
DESKTOP_W = Inches(5.8)


def log(m):
    print(m.encode("ascii", "replace").decode("ascii"))


def find_para(doc, text, exact=False):
    for p in doc.element.body.iter(qn("w:p")):
        para = Paragraph(p, doc)
        t = para.text.strip()
        if (exact and t == text) or (not exact and text in t):
            return para
    return None


def next_until_heading(start_p):
    el = start_p._p.getnext()
    while el is not None:
        if el.tag == qn("w:p"):
            p = Paragraph(el, start_p._parent)
            style = (p.style.name if p.style else "") or ""
            text = p.text.strip()
            if style.startswith("Heading") and text:
                break
            yield el, p
        else:
            yield el, None
        el = el.getnext()


def is_caption(text):
    t = (text or "").strip()
    return t.startswith("Implemented screen")


def clear_images_under(h):
    for el, p in list(next_until_heading(h)):
        if p is None:
            continue
        if el.findall(".//" + qn("w:drawing")) or is_caption(p.text):
            el.getparent().remove(el)


def insert_image_after(doc, anchor, img_path, width, caption):
    img_p = Paragraph(anchor._p.makeelement(qn("w:p"), {}), doc)
    anchor._p.addnext(img_p._p)
    img_p.alignment = 1
    img_p.add_run().add_picture(str(img_path), width=width)
    cap = Paragraph(img_p._p.makeelement(qn("w:p"), {}), doc)
    img_p._p.addnext(cap._p)
    cap.alignment = 1
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(10)
    return cap


def find_anchor(h):
    anchor = h
    for el, p in next_until_heading(h):
        if p is None:
            continue
        t = p.text.strip()
        if t.startswith("Related Use Case") or t.startswith("Related Use Cases"):
            return p
        if t.startswith("Platform:") or t.startswith("Primary Actor"):
            anchor = p
        if t.startswith("This screen") or t.startswith("This function") or t.startswith("This section"):
            break
    return anchor


def remove_no_ui_sentences(doc):
    patterns = [
        "This function is performed on the Customer Order Status screen. Therefore, no separate complex UI layout is required.",
        "This function is performed on the Cashier Order List and Payment Screen. Therefore, no separate UI layout is required.",
        "This function is performed on the Barista Order Board Screen. Therefore, no separate UI layout is required.",
    ]
    # also catch partial
    for p in list(doc.element.body.iter(qn("w:p"))):
        para = Paragraph(p, doc)
        t = para.text.strip()
        if not t:
            continue
        if "Therefore, no separate UI layout is required" in t or "Therefore, no separate complex UI layout is required" in t:
            # keep backend note and known-limitations
            if "backend services" in t.lower():
                continue
            if "payment timestamp" in t.lower() or "live monitoring" in t.lower():
                continue
            p.getparent().remove(p)
            log(f"removed no-UI sentence: {t[:70]}")


def fix_uc_diagrams(doc):
    """Cashier/Manager diagram images sit in empty Heading 3 paras — convert to Normal
    and delete empty Heading junk that causes page-break 'jump to top'."""
    kids = list(doc.element.body.iterchildren())
    # Find range from 1.3.2.4 to 1.4
    start_i = end_i = None
    for i, el in enumerate(kids):
        if el.tag != qn("w:p"):
            continue
        t = Paragraph(el, doc).text.strip()
        if t.startswith("1.3.2.4"):
            start_i = i
        if t.startswith("1.4 System Functionalities"):
            end_i = i
            break
    if start_i is None or end_i is None:
        log("!! cannot locate UC diagram range")
        return

    # Process paragraphs in range
    for i in range(start_i, end_i):
        el = kids[i]
        if el.tag != qn("w:p"):
            continue
        p = Paragraph(el, doc)
        style = (p.style.name if p.style else "") or ""
        text = p.text.strip()
        has_img = bool(el.findall(".//" + qn("w:drawing")))

        # Image wrongly styled as Heading -> Normal
        if has_img and style.startswith("Heading"):
            try:
                p.style = "Normal"
            except Exception:
                try:
                    p.style = "normal"
                except Exception:
                    pass
            log(f"UC diagram: image para style Heading->Normal near {i}")

        # Empty heading junk (no text, no image)
        if style.startswith("Heading") and not text and not has_img:
            # keep real section headings only — these empties are garbage
            el.set(qn("w:rsidR"), el.get(qn("w:rsidR")) or "00000000")  # noop touch
            # mark for deletion
            el.set("{http://mark}del", "1")

    # Second pass delete marked
    for el in list(doc.element.body.iterchildren()):
        if el.get("{http://mark}del") == "1":
            el.getparent().remove(el)
            log("removed empty heading junk")

    # Also clean empty Heading 5 between 1.3.2.4 note and 1.3.2.5
    kids = list(doc.element.body.iterchildren())
    in_gap = False
    for el in kids:
        if el.tag != qn("w:p"):
            continue
        p = Paragraph(el, doc)
        t = p.text.strip()
        st = (p.style.name if p.style else "") or ""
        if t.startswith("1.3.2.4"):
            in_gap = True
            continue
        if t.startswith("1.3.2.5"):
            in_gap = False
            continue
        if in_gap and st.startswith("Heading") and not t and not el.findall(".//" + qn("w:drawing")):
            el.getparent().remove(el)
            log("removed empty Heading in cashier-manager gap")


def replace_section_image(doc, heading, images):
    """images: list of (filename, caption)"""
    h = find_para(doc, heading, exact=True) or find_para(doc, heading)
    if h is None:
        log(f"!! missing {heading}")
        return None
    clear_images_under(h)
    anchor = find_anchor(h)
    # insert in reverse so order is correct after anchor
    for fname, caption in reversed(images):
        insert_image_after(doc, anchor, SHOTS / fname, DESKTOP_W, caption)
    log(f"set images under [{heading}] <- {[f for f,_ in images]}")
    return h


def add_section_333(doc):
    """Insert 3.3.3 Cook by Item Preparation after 3.3.2 block, before 3.4."""
    existing = find_para(doc, "3.3.3 Cook by Item Preparation", exact=True)
    if existing:
        replace_section_image(doc, "3.3.3 Cook by Item Preparation", [
            ("07-barista-item-prepare.png",
             "Implemented screen: Cook by Item mode — mark preparedQty per item line"),
        ])
        return

    h34 = find_para(doc, "3.4 Waiter Service Station", exact=True)
    if h34 is None:
        log("!! 3.4 missing")
        return

    def insert_after(anchor_el, text="", style=None, image=None, caption=None):
        p = Paragraph(anchor_el.makeelement(qn("w:p"), {}), doc)
        anchor_el.addnext(p._p)
        if style:
            try:
                p.style = style
            except Exception:
                pass
        if image:
            p.alignment = 1
            p.add_run().add_picture(str(image), width=DESKTOP_W)
            if caption:
                cap = Paragraph(p._p.makeelement(qn("w:p"), {}), doc)
                p._p.addnext(cap._p)
                cap.alignment = 1
                r = cap.add_run(caption)
                r.italic = True
                r.font.size = Pt(10)
                return cap._p
            return p._p
        if text:
            p.add_run(text)
        return p._p

    # heading before 3.4, then chain after heading
    h_el = h34._p.makeelement(qn("w:p"), {})
    h34._p.addprevious(h_el)
    h_para = Paragraph(h_el, doc)
    try:
        h_para.style = "Heading 4"
    except Exception:
        pass
    h_para.add_run("3.3.3 Cook by Item Preparation")
    cursor = h_el
    for text in [
        "Platform: Tablet Web / Desktop Web",
        "Primary Actor: Barista",
        "Related Use Case: UC-04 - Update Preparation Status",
    ]:
        cursor = insert_after(cursor, text)
    cursor = insert_after(
        cursor,
        image=SHOTS / "07-barista-item-prepare.png",
        caption="Implemented screen: Cook by Item mode — mark preparedQty per item line",
    )
    cursor = insert_after(cursor, "This screen allows the Barista to:")
    for text in [
        "Switch the preparation board to Cook by Item mode.",
        "See each ordered item line with preparedQty progress (for example 0/2).",
        "Mark individual item lines as prepared until the whole order becomes Ready.",
        "Trigger cup and recipe-ingredient deduction when preparation completes.",
    ]:
        cursor = insert_after(cursor, text)
    log("added section 3.3.3 Cook by Item Preparation")


def update_332_text(doc):
    """Clarify 3.3.2 is order-level Pending/Preparing/Ready status updates."""
    h = find_para(doc, "3.3.2 Update Item Status Function", exact=True)
    if not h:
        return
    # Optionally rename heading to Update Order Status Function — user said status update pending/ready
    # Keep ID 3.3.2 but adjust description bullets if needed
    for el, p in next_until_heading(h):
        if p and p.text.strip() == "Related Use Case: UC-04 - Update Preparation Status":
            # ensure there's a clarifying sentence after image later
            pass
    # Add clarifying bullets if missing
    texts = [p.text.strip() for _, p in next_until_heading(h) if p]
    if not any("Pending" in t and "Ready" in t and "Cook by Order" in t for t in texts):
        for el, p in next_until_heading(h):
            if p and p.text.strip().startswith("This function allows the Barista"):
                np = Paragraph(p._p.makeelement(qn("w:p"), {}), doc)
                p._p.addnext(np._p)
                np.add_run(
                    "Work in Cook by Order mode and move orders between Pending, Preparing, and Ready."
                )
                log("3.3.2: added order-status mode bullet")
                break


def main():
    shutil.copyfile(SRC, BACKUP)
    doc = Document(str(SRC))

    remove_no_ui_sentences(doc)
    fix_uc_diagrams(doc)

    # 3.2.1 cashier without split
    replace_section_image(doc, "3.2.1 POS Order List Screen", [
        ("10-cashier-unpaid.png",
         "Implemented screen: cashier unpaid/paid order list (no split-bill overlay)"),
    ])

    # 3.3.2 status update: pending + ready
    replace_section_image(doc, "3.3.2 Update Item Status Function", [
        ("06d-barista-status-update.png",
         "Implemented screen: barista order board — Pending queue (Cook by Order)"),
        ("06c-barista-ready.png",
         "Implemented screen: barista order board — Ready queue after status updates"),
    ])

    update_332_text(doc)
    add_section_333(doc)

    # Also refresh 3.3.1 board image (order mode) if present
    replace_section_image(doc, "3.3.1 KDS Order Board Screen", [
        ("06-barista-board.png",
         "Implemented screen: barista board by order (Pending / Preparing / Ready) with cup stock"),
    ])

    doc.save(str(SRC))
    log(f"saved {SRC}")


if __name__ == "__main__":
    main()
