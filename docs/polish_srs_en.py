# -*- coding: utf-8 -*-
"""Polish SRS_v2: complete UC coverage, English captions, re-embed EN screenshots,
unify table header colors, clean leftover bullets/dots.
"""
import copy
import re
import shutil
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table, _Row
from docx.text.paragraph import Paragraph

DOCS = Path(r"d:\STUDY\Tai_lieu\FPT_University\Semester\Summer_2026\SWP391\Documents")
SRC = DOCS / "Group2_1.SRS_v2.docx"
DST = DOCS / "Group2_1.SRS_v2.docx"
BACKUP = DOCS / "Group2_1.SRS_v2.before_en_polish.docx"
SHOTS = DOCS / "srs_screenshots"

MOBILE_W = Inches(2.9)
DESKTOP_W = Inches(5.8)
HEADER_FILL = "1F4E79"  # unified navy
HEADER_FONT = "FFFFFF"

SECTION_IMAGES = [
    ("3.1.1 Customer Web Menu Screen", "02-menu-customer.png", MOBILE_W,
     "Implemented screen: customer menu after table QR / tableCode lock"),
    ("3.1.2 Order Summary / Checkout Screen", "03-menu-cart.png", MOBILE_W,
     "Implemented screen: cart with quantity controls, remove item, total, and order note"),
    ("3.1.3 Table Order Tracking Function", "04-order-status.png", MOBILE_W,
     "Implemented screen: order status for the locked table / current guest session"),
    ("3.2.1 POS Order List Screen", "10-cashier-unpaid.png", DESKTOP_W,
     "Implemented screen: cashier unpaid/paid tabs (no order search box)"),
    ("3.2.2 Payment Processing Function", "11-cashier-split.png", DESKTOP_W,
     "Implemented screen: split-bill dialog on cashier (payment is per Served order)"),
    ("3.2.3 Counter Order Function", "12-counter-order.png", DESKTOP_W,
     "Implemented screen: counter order (Cashier/Admin selects table and items)"),
    ("3.3.1 KDS Order Board Screen", "06-barista-board.png", DESKTOP_W,
     "Implemented screen: barista board by order (Pending / Preparing / Ready) with cup stock"),
    ("3.3.2 Update Item Status Function", "07-barista-item-prepare.png", DESKTOP_W,
     "Implemented screen: Cook by Item mode — mark preparedQty per item line"),
    ("3.4.1 Wait Station / Table Layout Screen", "08-runner-station.png", DESKTOP_W,
     "Implemented screen: waiter/runner station — serve, clear table, table map"),
    ("3.4.2 Move / Merge Table Function", "09-table-transfer.png", DESKTOP_W,
     "Implemented screen: table transfer (move all open orders to the target table)"),
    ("3.5.1 Admin Dashboard Screen", "14-admin-dashboard.png", DESKTOP_W,
     "Implemented screen: admin dashboard — revenue, best sellers, low-stock alert, table map"),
    ("3.5.2 Product and Menu Management Screen", "15-admin-menu.png", DESKTOP_W,
     "Implemented screen: menu CRUD, recipes, Download Template / Import from Excel"),
    ("3.5.3 Table and QR Management Screen", "16-admin-tables.png", DESKTOP_W,
     "Implemented screen: tables and QR management"),
    ("3.5.4 Reports Dashboard Screen", "14b-admin-reports.png", DESKTOP_W,
     "Implemented screen: revenue report filters on the dashboard (day/week/month/year)"),
    ("3.5.5 Role Access and Demo Staff Accounts", "13-admin-pin-gate.png", DESKTOP_W,
     "Implemented screen: admin PIN gate (8888) before unlocking the dashboard"),
    ("3.5.6 System Logs Screen", "20-system-logs.png", DESKTOP_W,
     "Implemented screen: system logs filtered by actor/role (system-logs.jsp)"),
    ("3.6.1 Staff PIN Login Screen", "05-staff-login.png", DESKTOP_W,
     "Implemented screen: staff PIN login by role (Barista 1111 / Cashier 2222 / Waiter-Runner 3333)"),
    ("3.8.1 Staff Management Screen", "17-admin-staff.png", DESKTOP_W,
     "Implemented screen: staff CRUD and weekly shift board (overlap prevention)"),
    ("3.8.2 Monthly Payroll Section", "18-admin-payroll.png", DESKTOP_W,
     "Implemented screen: monthly payroll (Evening=5h, other shifts=6h; only completed shifts)"),
    ("3.9.1 Inventory Screen", "19-admin-inventory.png", DESKTOP_W,
     "Implemented screen: ingredient inventory (stock, min stock, low-stock warnings)"),
]


def log(msg):
    print(msg.encode("ascii", "replace").decode("ascii"))


def all_tables(doc):
    return [Table(t, doc) for t in doc.element.body.iter(qn("w:tbl"))]


def find_para(doc, text, exact=False):
    for p in doc.element.body.iter(qn("w:p")):
        para = Paragraph(p, doc)
        t = para.text.strip()
        if (exact and t == text) or (not exact and text in t):
            return para
    return None


def set_cell(cell, value):
    ps = cell.paragraphs
    for extra in ps[1:]:
        extra._p.getparent().remove(extra._p)
    p = ps[0]
    for r in list(p._p.findall(qn("w:r"))):
        p._p.remove(r)
    lines = value if isinstance(value, list) else [value]
    for i, line in enumerate(lines):
        run = p.add_run(line)
        if i < len(lines) - 1:
            run._r.append(run._r.makeelement(qn("w:br"), {}))


def set_para_text(para, text):
    for r in list(para._p.findall(qn("w:r"))):
        para._p.remove(r)
    para.add_run(text)


def clone_row(table, template_idx, texts, insert_after_idx=None):
    tr = copy.deepcopy(table.rows[template_idx]._tr)
    if insert_after_idx is None:
        table._tbl.append(tr)
    else:
        table.rows[insert_after_idx]._tr.addnext(tr)
    row = _Row(tr, table)
    for cell, text in zip(row.cells, texts):
        set_cell(cell, text)
    return row


def next_until_heading(start_p):
    el = start_p._p.getnext()
    while el is not None:
        if el.tag == qn("w:p"):
            p = Paragraph(el, start_p._parent)
            style = (p.style.name if p.style else "") or ""
            text = p.text.strip()
            if style.startswith("Heading") and text:
                break
            yield el, p
        else:
            yield el, None
        el = el.getnext()


def is_caption(text):
    t = (text or "").strip()
    return t.startswith("Implemented screen") or t.startswith("Màn hình thực tế")


# ---------- catalog / coverage / auth ----------

def update_catalog(doc, tables):
    cat = [t for t in tables if len(t.columns) == 4 and "Use Case Description" in t.rows[0].cells[3].text][0]

    def uc_name(r):
        return r.cells[1].text.strip().split("\n")[0].strip()

    # Rename ambiguous menu UC
    for r in cat.rows:
        if uc_name(r).startswith("Manage Inventory") and "menu items" in r.cells[3].text.lower():
            set_cell(r.cells[1], "Manage Menu Products\n «include» Record System Log")
            set_cell(r.cells[2], "Menu Configuration")
            set_cell(r.cells[3],
                     "Manager creates, edits, activates/deactivates menu products with bilingual names, "
                     "category, price, image, optional sizes, and recipe ingredients; each change is logged.")
            log("catalog: renamed Manage Inventory -> Manage Menu Products")

    # Enrich Prepare Item Line
    for r in cat.rows:
        if uc_name(r) == "Prepare Item Line":
            set_cell(r.cells[1], "Prepare Item Line\n «extend» of View Order Queue")
            set_cell(r.cells[3],
                     "Barista switches to Cook by Item mode and marks each item line prepared "
                     "(preparedQty). When every line is fully prepared the order becomes Ready "
                     "automatically; cups and recipe ingredients are deducted.")
            log("catalog: enriched Prepare Item Line (Cook by Item)")

    # Enrich Manage Staff & Shifts
    for r in cat.rows:
        if uc_name(r).startswith("Manage Staff & Shifts"):
            set_cell(r.cells[3],
                     "Manager manages staff records (name, role, personal PIN, active status) and assigns "
                     "Morning/Afternoon/Evening shifts without overlaps (SHIFT_OVERLAP); each change is logged.")

    # Insert missing UCs if absent
    names = {uc_name(r) for r in cat.rows}
    inserts = []
    if not any(n.startswith("Switch Cook Mode") for n in names):
        inserts.append(("Mark Order Ready",
                        ["0", "Switch Cook Mode (Order / Item)\n «extend» of View Order Queue",
                         "Order Preparation",
                         "Barista switches the preparation board between Cook by Order and Cook by Item modes."]))
    if not any(n.startswith("Update Cup Stock") for n in names):
        inserts.append(("View Revenue Dashboard",
                        ["0", "Update Cup Stock\n «include» Record System Log", "Configuration",
                         "Manager updates the available cup count (set/add/subtract); Barista sees the "
                         "updated cup stock on the preparation board."]))
    if not any(n.startswith("View Low-Stock") for n in names):
        inserts.append(("View Revenue Dashboard",
                        ["0", "View Low-Stock Warning\n «extend» of View Revenue Dashboard", "Inventory Monitoring",
                         "Dashboard shows ingredients below minimum stock and links to Inventory management."]))

    for after_name, texts in inserts:
        for i, r in enumerate(cat.rows):
            if uc_name(r).startswith(after_name.split("\n")[0]):
                clone_row(cat, i, texts, insert_after_idx=i)
                log(f"catalog: added {texts[1].split(chr(10))[0]}")
                break

    # Renumber
    n = 0
    for r in cat.rows:
        if r.cells[0].text.strip().isdigit() or r.cells[0].text.strip() == "0":
            n += 1
            set_cell(r.cells[0], str(n))
    log(f"catalog renumbered 1..{n}")


def update_auth_table(doc, tables):
    auth = [t for t in tables if t.rows[0].cells[0].text.strip() == "Screen"
            and "Manager" in t.rows[0].cells[1].text][0]
    existing = {r.cells[0].text.strip() for r in auth.rows}
    extras = [
        ["Staff Management (Shifts & Payroll)", "X", "", "", "", ""],
        ["Ingredient Inventory", "X", "", "", "", ""],
        ["System Logs", "X", "", "", "", ""],
        ["Counter Order", "X", "X", "", "", ""],
        ["Table Transfer", "X", "X", "", "X", ""],
    ]
    last = len(auth.rows) - 1
    for row in extras:
        if row[0] not in existing:
            clone_row(auth, last, row)
            log(f"auth: added {row[0]}")


def ensure_fr_bullets(doc):
    """Ensure key feature bullets exist in FR sections."""
    # 3.3.1 cook modes
    h = find_para(doc, "3.3.1 KDS Order Board Screen", exact=True)
    if h:
        texts = [p.text.strip() for _, p in next_until_heading(h) if p]
        if not any("Cook by Item" in t or "cook by item" in t.lower() for t in texts):
            # insert after "This screen allows the Barista to:"
            for el, p in next_until_heading(h):
                if p and p.text.strip().startswith("This screen allows the Barista"):
                    for bullet in [
                        "Switch between Cook by Order and Cook by Item modes.",
                        "In Cook by Item mode, mark each item line prepared and track preparedQty.",
                    ]:
                        np = Paragraph(p._p.makeelement(qn("w:p"), {}), doc)
                        p._p.addnext(np._p)
                        np.add_run(bullet)
                        try:
                            np.style = p.style
                        except Exception:
                            pass
                        p = np
                    log("3.3.1: added Cook by Item bullets")
                    break

    # 3.5.2 excel import
    h = find_para(doc, "3.5.2 Product and Menu Management Screen", exact=True)
    if h:
        texts = [p.text.strip() for _, p in next_until_heading(h) if p]
        if not any("Excel" in t or "import" in t.lower() for t in texts):
            for el, p in next_until_heading(h):
                if p and p.text.strip().startswith("This screen allows the Manager"):
                    for bullet in [
                        "Download a generated Excel import template.",
                        "Import menu products in bulk from a filled Excel file (invalid rows are skipped).",
                        "Define ingredient recipes for each menu item used when preparation completes.",
                    ]:
                        np = Paragraph(p._p.makeelement(qn("w:p"), {}), doc)
                        p._p.addnext(np._p)
                        np.add_run(bullet)
                        p = np
                    log("3.5.2: added Excel/recipe bullets")
                    break

    # 3.5.1 low stock
    h = find_para(doc, "3.5.1 Admin Dashboard Screen", exact=True)
    if h:
        texts = [p.text.strip() for _, p in next_until_heading(h) if p]
        if not any("low-stock" in t.lower() or "low stock" in t.lower() for t in texts):
            for el, p in next_until_heading(h):
                if p and p.text.strip().startswith("This screen allows the Manager"):
                    np = Paragraph(p._p.makeelement(qn("w:p"), {}), doc)
                    p._p.addnext(np._p)
                    np.add_run("View low-stock ingredient warnings and open Inventory management.")
                    log("3.5.1: added low-stock bullet")
                    break


# ---------- images ----------

def replace_section_images(doc):
    for heading, img, width, caption in SECTION_IMAGES:
        h = find_para(doc, heading, exact=True) or find_para(doc, heading)
        if h is None:
            log(f"!! missing heading {heading}")
            continue
        # remove existing drawings + captions under section
        for el, p in list(next_until_heading(h)):
            if p is None:
                continue
            if el.findall(".//" + qn("w:drawing")) or is_caption(p.text):
                el.getparent().remove(el)
        # find insert anchor
        anchor = h
        for el, p in next_until_heading(h):
            if p is None:
                continue
            t = p.text.strip()
            if t.startswith("Related Use Case") or t.startswith("Related Use Cases"):
                anchor = p
                break
            if t.startswith("Platform:") or t.startswith("Primary Actor"):
                anchor = p
            if t.startswith("This screen") or t.startswith("This function") or t.startswith("This section"):
                break
        path = SHOTS / img
        if not path.exists():
            log(f"!! missing image {img}")
            continue
        img_p = Paragraph(anchor._p.makeelement(qn("w:p"), {}), doc)
        anchor._p.addnext(img_p._p)
        img_p.alignment = 1
        img_p.add_run().add_picture(str(path), width=width)
        cap = Paragraph(img_p._p.makeelement(qn("w:p"), {}), doc)
        img_p._p.addnext(cap._p)
        cap.alignment = 1
        r = cap.add_run(caption)
        r.italic = True
        r.font.size = Pt(10)
        log(f"image [{heading[:30]}] <- {img}")


# ---------- English leftovers + diagram notes ----------

def translate_and_clean_text(doc):
    replacements = [
        ("Màn hình thực tế:", "Implemented screen:"),
        ("Note: bill splitting is performed by the Cashier (or Admin) only; the Barista has no split-bill "
         "function in the implemented system.",
         "Note: bill splitting is performed by the Cashier (or Admin) only; the Barista has no split-bill "
         "function in the implemented system."),
        ("Note: the diagram above shows the core entities. The implementation additionally uses the "
         "Staff, Shifts, Inventory, and RecipeItems tables described below (13 tables in total).",
         "Note: the diagram above shows the core entities. The implementation additionally uses the "
         "Staff, Shifts, Inventory, and RecipeItems tables described below (13 tables in total)."),
    ]
    # Translate any remaining Vietnamese captions wholesale via known map
    vi_to_en = {
        # handled by image replace captions
    }
    for p in doc.element.body.iter(qn("w:p")):
        para = Paragraph(p, doc)
        t = para.text
        if not t.strip():
            continue
        # remove lone bullet leftovers
        if t.strip() in (".", "-", "•", "·"):
            p.getparent().remove(p)
            log("removed leftover bullet/dot")
            continue
        # strip leading orphan punctuation like ". text"
        if re.match(r"^[.\-•·]\s+\S", t.strip()):
            set_para_text(para, re.sub(r"^[.\-•·]\s+", "", t.strip()))
            log("cleaned leading bullet char")
        # Vietnamese caption leftover (if any after image replace)
        if "Màn hình thực tế" in t:
            set_para_text(para, t.replace("Màn hình thực tế", "Implemented screen"))
            log("translated leftover VI caption")


def clean_multiline_bullets_in_specs(doc):
    """Split jammed alternative-flow lines that lack separators into clean sentences.
    Only touch paragraphs that still contain 'A1 -' jammed with 'A2 -' on one visual block
    already handled in prior revision — skip if fine.
    """
    for p in doc.element.body.iter(qn("w:p")):
        para = Paragraph(p, doc)
        t = para.text.strip()
        # empty list items that are only whitespace + dash in runs
        if t == "-" or t == ".":
            p.getparent().remove(p)


# ---------- table colors ----------

def unify_table_headers(doc):
    tables = all_tables(doc)
    for ti, table in enumerate(tables):
        if not table.rows:
            continue
        for cell in table.rows[0].cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            # remove old shd
            for old in list(tcPr.findall(qn("w:shd"))):
                tcPr.remove(old)
            shd = tcPr.makeelement(qn("w:shd"), {
                qn("w:val"): "clear",
                qn("w:color"): "auto",
                qn("w:fill"): HEADER_FILL,
            })
            tcPr.append(shd)
            # white bold text in header
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.bold = True
        # lightly tint alt body rows for readability (optional skip)
    log(f"unified header color #{HEADER_FILL} on {len(tables)} tables")


def update_record_of_changes(doc, tables):
    roc = tables[0]
    clone_row(roc, len(roc.rows) - 1, [
        "21/7", "A, M", "Group 2",
        "Completed use-case coverage for Cook-by-Item, staff/shifts/payroll, ingredient inventory, "
        "recipes, and Excel import; expanded screen authorization; replaced FR screenshots with English "
        "UI captures; translated captions to English; unified table header styling; cleaned leftover bullets."
    ])


def main():
    shutil.copyfile(SRC, BACKUP)
    doc = Document(str(SRC))
    tables = all_tables(doc)

    update_catalog(doc, tables)
    update_auth_table(doc, tables)
    ensure_fr_bullets(doc)
    replace_section_images(doc)
    translate_and_clean_text(doc)
    clean_multiline_bullets_in_specs(doc)
    unify_table_headers(doc)
    update_record_of_changes(doc, tables)

    # TOC refresh flag
    settings = doc.settings.element
    if settings.find(qn("w:updateFields")) is None:
        settings.append(settings.makeelement(qn("w:updateFields"), {qn("w:val"): "true"}))

    doc.save(str(DST))
    log(f"saved {DST}")


if __name__ == "__main__":
    main()
